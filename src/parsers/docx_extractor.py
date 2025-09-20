"""
DOCX Text Extraction Module
Handles extraction of text from DOCX files using python-docx and docx2txt
"""

from docx import Document
import docx2txt
from typing import Dict, Any, List
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

class DOCXExtractor:
    """Extract text from DOCX files using multiple methods for robustness"""
    
    def __init__(self):
        self.methods = ['python-docx', 'docx2txt']
    
    def extract_text_python_docx(self, file_path: str) -> str:
        """Extract text using python-docx library"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"python-docx extraction failed for {file_path}: {str(e)}")
            return ""
    
    def extract_text_docx2txt(self, file_path: str) -> str:
        """Extract text using docx2txt library"""
        try:
            text = docx2txt.process(file_path)
            return text if text else ""
        
        except Exception as e:
            logger.error(f"docx2txt extraction failed for {file_path}: {str(e)}")
            return ""
    
    def extract_text(self, file_path: str, method: str = 'auto') -> Dict[str, Any]:
        """
        Extract text from DOCX with fallback methods
        
        Args:
            file_path: Path to DOCX file
            method: Extraction method ('auto', 'python-docx', 'docx2txt')
        
        Returns:
            Dictionary with extracted text and metadata
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        result = {
            'text': '',
            'method_used': '',
            'success': False,
            'error': None
        }
        
        try:
            if method == 'auto':
                # Try python-docx first, then docx2txt as fallback
                text = self.extract_text_python_docx(file_path)
                if text and len(text.strip()) > 10:
                    result['text'] = text
                    result['method_used'] = 'python-docx'
                    result['success'] = True
                else:
                    text = self.extract_text_docx2txt(file_path)
                    if text:
                        result['text'] = text
                        result['method_used'] = 'docx2txt'
                        result['success'] = True
            
            elif method == 'python-docx':
                text = self.extract_text_python_docx(file_path)
                result['text'] = text
                result['method_used'] = 'python-docx'
                result['success'] = bool(text)
            
            elif method == 'docx2txt':
                text = self.extract_text_docx2txt(file_path)
                result['text'] = text
                result['method_used'] = 'docx2txt'
                result['success'] = bool(text)
            
            else:
                raise ValueError(f"Unknown extraction method: {method}")
                
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        
        return result
    
    def extract_structured_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract structured content from DOCX including formatting and metadata
        
        Returns:
            Dictionary with text, tables, formatting, and document metadata
        """
        result = {
            'text': '',
            'tables': [],
            'headings': [],
            'metadata': {},
            'success': False,
            'error': None
        }
        
        try:
            doc = Document(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'category': core_props.category or '',
                'comments': core_props.comments or ''
            }
            
            # Extract paragraphs with formatting
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        result['headings'].append({
                            'text': paragraph.text,
                            'level': paragraph.style.name,
                            'position': len(text_parts)
                        })
                    text_parts.append(paragraph.text)
            
            result['text'] = "\n".join(text_parts)
            
            # Extract tables with structure
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                result['tables'].append({
                    'table_id': table_num + 1,
                    'data': table_data,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0
                })
            
            result['success'] = True
        
        except Exception as e:
            result['error'] = str(e)
            logger.error(f"Structured DOCX extraction failed for {file_path}: {str(e)}")
            
            # Fallback to simple text extraction
            fallback = self.extract_text(file_path)
            result['text'] = fallback['text']
            result['success'] = fallback['success']
        
        return result